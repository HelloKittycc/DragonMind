from io import BytesIO

from src.modules.knowledge.parsers.base import KnowledgeFileParser, ParserError


class DocxKnowledgeParser(KnowledgeFileParser):
    supported_extensions = {".docx"}

    def extract_text(self, content: bytes, filename: str) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise ParserError("DOCX parser dependency is not installed") from exc

        try:
            document = Document(BytesIO(content))
        except Exception as exc:
            raise ParserError("DOCX file could not be parsed") from exc

        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table_index, table in enumerate(document.tables, start=1):
            rows: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                parts.append(f"[Table {table_index}]\n" + "\n".join(rows))
        return "\n\n".join(parts)
